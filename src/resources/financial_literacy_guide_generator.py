"""
Financial Literacy Guide Generator for Small Businesses

This script generates comprehensive financial literacy guides for small businesses
in Word format. It produces modular sections that can be downloaded individually
or as a complete guide.

Requirements:
- python-docx
- python-docx-template (for more advanced templates)

Install with: pip install python-docx python-docx-template
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import additional sections
from financial_literacy_sections import get_additional_sections

# Configuration
COMPANY_NAME = "Clarity Impact Finance"
OUTPUT_DIR = "guides_output"
BRAND_COLOR_GREEN = RGBColor(0, 128, 96)  # RGB values for brand green
BRAND_COLOR_ORANGE = RGBColor(242, 101, 34)  # RGB values for brand orange

# Guide structure with content
GUIDE_STRUCTURE = {
    "Introduction": {
        "content": [
            {
                "type": "paragraph",
                "text": "Financial literacy is fundamental to small business success. Without a solid understanding of financial concepts and practices, even the most promising businesses can struggle to survive and thrive."
            },
            {
                "type": "paragraph",
                "text": "This comprehensive guide is designed to help small business owners and entrepreneurs develop the financial knowledge and skills necessary to qualify for and effectively manage CDFI (Community Development Financial Institution) financing."
            },
            {
                "type": "paragraph",
                "text": "How to use this guide: You can read this guide from start to finish for a complete financial education, or you can focus on specific sections that address your immediate needs. Each section stands on its own while contributing to a comprehensive understanding of small business finances."
            },
            {
                "type": "callout",
                "title": "Why Financial Literacy Matters",
                "text": "Studies show that businesses with strong financial management are 30% more likely to survive their first five years and twice as likely to secure financing when needed."
            }
        ],
        "action_steps": [
            "Assess your current financial knowledge using the self-evaluation tool",
            "Identify which sections of this guide address your most pressing financial challenges",
            "Schedule dedicated time to work through the material and implement recommendations"
        ],
        "success_story": {
            "title": "From Financial Confusion to Clarity",
            "text": "Maria's bakery had been operating for three years when she approached a CDFI for growth capital. Initially rejected due to disorganized finances, Maria used this guide to restructure her financial reporting. Six months later, she secured a $75,000 loan to expand her business, which has since doubled its revenue."
        },
        "resources": [
            "Financial Literacy Self-Assessment Tool",
            "Business Financial Health Checklist",
            "Guide to Financial Terms and Concepts"
        ]
    },
    
    "Financial Fundamentals": {
        "content": [
            {
                "type": "paragraph",
                "text": "Understanding basic accounting concepts is essential for making informed business decisions. This section covers the fundamental building blocks of financial literacy."
            },
            {
                "type": "subheading",
                "text": "Basic Accounting Concepts"
            },
            {
                "type": "bullet_list",
                "items": [
                    "Assets: Resources owned by your business that have economic value",
                    "Liabilities: Debts or obligations your business owes to others",
                    "Equity: The residual interest in the assets after deducting liabilities",
                    "Revenue: Income generated from business activities",
                    "Expenses: Costs incurred to run your business"
                ]
            },
            {
                "type": "subheading",
                "text": "Understanding Financial Statements"
            },
            {
                "type": "paragraph",
                "text": "Financial statements provide a snapshot of your business's financial health and performance."
            },
            {
                "type": "bullet_list",
                "items": [
                    "Income Statement: Shows revenue, expenses, and profit over a specific period",
                    "Balance Sheet: Displays assets, liabilities, and equity at a point in time",
                    "Cash Flow Statement: Tracks the flow of cash in and out of your business"
                ]
            },
            {
                "type": "subheading",
                "text": "Key Financial Ratios"
            },
            {
                "type": "table",
                "headers": ["Ratio", "Formula", "What It Tells You"],
                "rows": [
                    ["Current Ratio", "Current Assets ÷ Current Liabilities", "Ability to pay short-term obligations"],
                    ["Debt-to-Equity", "Total Liabilities ÷ Total Equity", "Financial leverage and risk"],
                    ["Profit Margin", "Net Income ÷ Revenue", "Profitability efficiency"],
                    ["Inventory Turnover", "Cost of Goods Sold ÷ Average Inventory", "Sales and inventory management efficiency"]
                ]
            },
            {
                "type": "subheading",
                "text": "Cash Flow vs. Profit"
            },
            {
                "type": "paragraph",
                "text": "Many business owners confuse profit with cash flow. A business can be profitable on paper but still run out of cash. Understanding this distinction is crucial for survival."
            },
            {
                "type": "callout",
                "title": "Cash Flow vs. Profit Example",
                "text": "A business invoices $10,000 for services in January (counted as revenue) but doesn't receive payment until March. Meanwhile, it must pay employees and suppliers in January and February. This business may show a profit for January but face a cash shortfall."
            }
        ],
        "action_steps": [
            "Create or review your chart of accounts to ensure proper categorization",
            "Generate and review all three financial statements monthly",
            "Calculate key financial ratios quarterly and track trends",
            "Develop a cash flow projection for the next six months"
        ],
        "warning_signs": [
            "Consistent difficulty paying bills on time despite showing profits",
            "Rapidly increasing sales without corresponding cash increase",
            "Deteriorating financial ratios over multiple periods",
            "Inability to explain variations in financial statements"
        ],
        "resources": [
            "Financial Statement Template Pack",
            "Ratio Analysis Calculator",
            "Cash Flow Projection Tool"
        ]
    }
}

# More sections will be added dynamically by the generate_additional_sections function

def create_document_with_styles():
    """Create a document with predefined styles for consistent formatting."""
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = COMPANY_NAME
    core_properties.title = "Small Business Financial Literacy Guide"
    core_properties.created = datetime.now()
    
    # Define styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Guide Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = BRAND_COLOR_GREEN
    
    # Heading 1 style
    h1_style = styles.add_style('Guide Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = BRAND_COLOR_GREEN
    
    # Heading 2 style
    h2_style = styles.add_style('Guide Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = BRAND_COLOR_ORANGE
    
    # Normal text style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Callout style
    callout_style = styles.add_style('Guide Callout', WD_STYLE_TYPE.PARAGRAPH)
    callout_font = callout_style.font
    callout_font.name = 'Calibri'
    callout_font.size = Pt(11)
    callout_font.italic = True
    
    # Success story style
    success_style = styles.add_style('Guide Success Story', WD_STYLE_TYPE.PARAGRAPH)
    success_font = success_style.font
    success_font.name = 'Calibri'
    success_font.size = Pt(11)
    success_font.color.rgb = BRAND_COLOR_GREEN
    
    # Warning style
    warning_style = styles.add_style('Guide Warning', WD_STYLE_TYPE.PARAGRAPH)
    warning_font = warning_style.font
    warning_font.name = 'Calibri'
    warning_font.size = Pt(11)
    warning_font.color.rgb = RGBColor(192, 0, 0)  # Dark red for warnings
    
    return doc

def add_section_to_document(doc, section_title, section_data):
    """Add a complete section to the document with all its components."""
    # Add section heading
    heading = doc.add_paragraph(section_title, style='Guide Heading 1')
    
    # Add content
    for item in section_data.get('content', []):
        if item['type'] == 'paragraph':
            p = doc.add_paragraph(item['text'], style='Normal')
        
        elif item['type'] == 'subheading':
            p = doc.add_paragraph(item['text'], style='Guide Heading 2')
        
        elif item['type'] == 'bullet_list':
            for bullet_item in item['items']:
                p = doc.add_paragraph(bullet_item, style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.25)
        
        elif item['type'] == 'callout':
            doc.add_paragraph(item['title'], style='Guide Callout').bold = True
            p = doc.add_paragraph(item['text'], style='Guide Callout')
            # Add a light gray shading to callout paragraphs
            p = doc.add_paragraph()  # Add empty paragraph after callout
        
        elif item['type'] == 'table':
            table = doc.add_table(rows=1, cols=len(item['headers']))
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(item['headers']):
                header_cells[i].text = header
                # Make headers bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add rows
            for row_data in item['rows']:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data
            
            doc.add_paragraph()  # Add space after table
    
    # Add Action Steps
    if section_data.get('action_steps'):
        doc.add_paragraph('Action Steps', style='Guide Heading 2')
        for step in section_data['action_steps']:
            p = doc.add_paragraph(step, style='List Bullet')
            p.style = 'List Bullet'
    
    # Add Success Story
    if section_data.get('success_story'):
        story = section_data['success_story']
        doc.add_paragraph('Success Story: ' + story['title'], style='Guide Success Story').bold = True
        doc.add_paragraph(story['text'], style='Guide Success Story')
    
    # Add Warning Signs
    if section_data.get('warning_signs'):
        doc.add_paragraph('Warning Signs to Watch For', style='Guide Warning').bold = True
        for warning in section_data['warning_signs']:
            p = doc.add_paragraph(warning, style='List Bullet')
            # Make this a red bullet list
            for run in p.runs:
                run.font.color.rgb = RGBColor(192, 0, 0)
    
    # Add Resources
    if section_data.get('resources'):
        doc.add_paragraph('Additional Resources', style='Guide Heading 2')
        for resource in section_data['resources']:
            doc.add_paragraph('• ' + resource, style='Normal')
    
    # Add page break after each section
    doc.add_page_break()

def generate_additional_sections():
    """Generate content for additional guide sections."""
    # Import the additional sections from the separate file
    additional_sections = get_additional_sections()
    
    # Merge with the main guide structure
    GUIDE_STRUCTURE.update(additional_sections)

def generate_full_guide():
    """Generate the complete financial literacy guide."""
    # Ensure output directory exists
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    # Create document with styles
    doc = create_document_with_styles()
    
    # Add title page
    title = doc.add_paragraph("Small Business Financial Literacy Guide", style='Guide Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("A Comprehensive Resource for CDFI Clients")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Provided by {COMPANY_NAME}")
    doc.add_paragraph(f"Created: {datetime.now().strftime('%B %Y')}")
    doc.add_paragraph()
    doc.add_paragraph("This guide provides small business owners with essential financial knowledge needed to qualify for and effectively manage CDFI financing. Each section includes practical advice, worksheets, and action steps.")
    
    # Add page break after title page
    doc.add_page_break()
    
    # Add table of contents placeholder
    doc.add_paragraph("Table of Contents", style='Guide Heading 1')
    for section_title in GUIDE_STRUCTURE.keys():
        doc.add_paragraph(section_title)
    
    doc.add_page_break()
    
    # Add each section
    for section_title, section_data in GUIDE_STRUCTURE.items():
        add_section_to_document(doc, section_title, section_data)
    
    # Save the document
    full_guide_path = os.path.join(OUTPUT_DIR, "Small_Business_Financial_Literacy_Complete_Guide.docx")
    doc.save(full_guide_path)
    print(f"Complete guide saved to: {full_guide_path}")
    
    return full_guide_path

def generate_individual_sections():
    """Generate individual section documents that can be downloaded separately."""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    section_paths = []
    
    for section_title, section_data in GUIDE_STRUCTURE.items():
        # Create a new document for each section
        doc = create_document_with_styles()
        
        # Add section title as document title
        title = doc.add_paragraph(section_title, style='Guide Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph("Small Business Financial Literacy Guide")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Provided by {COMPANY_NAME}")
        doc.add_paragraph(f"Created: {datetime.now().strftime('%B %Y')}")
        
        doc.add_page_break()
        
        # Add section content
        add_section_to_document(doc, section_title, section_data)
        
        # Save individual section
        safe_title = section_title.replace(" ", "_").replace("/", "_")
        section_path = os.path.join(OUTPUT_DIR, f"Financial_Literacy_{safe_title}.docx")
        doc.save(section_path)
        section_paths.append(section_path)
        
        print(f"Section '{section_title}' saved to: {section_path}")
    
    return section_paths

def main():
    """Main function to run the guide generator."""
    print(f"Generating Small Business Financial Literacy Guide for {COMPANY_NAME}")
    
    # Generate additional sections
    generate_additional_sections()
    
    # Generate full guide
    full_guide_path = generate_full_guide()
    
    # Generate individual sections
    section_paths = generate_individual_sections()
    
    print("\nGuide generation complete!")
    print(f"Full guide: {full_guide_path}")
    print(f"Individual sections: {len(section_paths)} files created in {OUTPUT_DIR} directory")

if __name__ == "__main__":
    main()
